# -*- coding: utf-8 -*-
"""
Created on Wed Mar 23 22:52:48 2022

@author: sesha (Ani T)
"""
###### LAST EDITED: JOHN LEHAN 04/17/2023 #######
#######################################################



import csv
import time
from docx import Document
start_time = time.time()

header_processed = False # Used later to skip column headers

# Splitter designed for Word Doc w/ default font/size; could also use the page break alternative seen above
splitter = '------------------------------------------------------------------------------------------------------------------------------------------'

# Opening responses CSV file to be read (special method to ensure commas and multi paragraph responses can be used)
# CHANGE CSV FILENAME HERE TO OPEN CORRECT DOCUMENT
# Make sure this document is in the same file location as the code
with open('BanquetLetters.csv', newline = '', encoding='utf-8') as responses:
   # Reader to turn each line into list (special method to ensure commas and multi paragraph responses can be used)
   reader = csv.reader(responses, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
   
   # Loop for each individual line in responses CSV
   for r in reader:
       # Making sure to skip first row of column headers
       if header_processed:
           # r[2] based on specific Google Form; number may need to be adjusted for future use
           # The recipient defines the filename, so it's incredibly important that the recipient question on the Google Form is a dropdown, to avoid separate files getting generated mispellings (e.g someone types Jhon Lehan rather than John Lehan)
           filename = r[2] + ' Banquet Letters' + '.docx'
           # Writing formatted text to specific person's letters file
           try:
               recipient = Document(filename)
               # recipient.add_page_break() # Alternative to the splitter, allows separate pages
           # If the document does not already exist, open the templated document
           except:
               recipient = Document("BLTRLetterTemplate.docx") # Make sure this document is in the same file location as the code
               recipient.add_paragraph(r[2] + "'s Letters")
               recipient.add_paragraph(splitter)
               
           sender = r[4] # r[4] based on specific Google Form; number may need to be adjusted for future use
           letter = r[3] # r[3] based on specific Google Form; number may need to be adjusted for future use
       
           # This appends to the existing document; don't run the code more than once w/o deleting all the generated files first
           address = recipient.add_paragraph("From: " + sender) 
           body = recipient.add_paragraph(letter)
           split = recipient.add_paragraph(splitter)
           recipient.save(filename)    
       header_processed = True

print("--- Run Time = %s seconds ---" % (time.time() - start_time))

